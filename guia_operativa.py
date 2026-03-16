"""
Genera la Guía Operativa de Roker Nexus en formato .docx
Ejecutar: python guia_operativa.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ─────────────────────────────────────────────────
C_TITULO    = RGBColor(0x0A, 0x84, 0xFF)   # azul nexus
C_NARANJA   = RGBColor(0xFF, 0x9F, 0x0A)   # amber/warning
C_VERDE     = RGBColor(0x32, 0xD7, 0x4B)   # verde ok
C_ROJO      = RGBColor(0xFF, 0x3B, 0x30)   # rojo error
C_GRIS      = RGBColor(0x6E, 0x6E, 0x73)   # gris texto2
C_NEGRO     = RGBColor(0x1C, 0x1C, 0x1E)   # negro
C_FONDO_TBL = RGBColor(0xF2, 0xF2, 0xF7)   # fondo tabla header


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _run(para, text, bold=False, italic=False, color=None, size=None, font="Arial"):
    r = para.add_run(text)
    r.font.name = font
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return r


def _heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = color or C_TITULO
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after  = Pt(6)
    return h


def _para(doc, text="", indent=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    if text:
        _run(p, text, size=10.5)
    return p


def _step_box(doc, numero, titulo, pasos, nota=None, advertencia=None):
    """Caja de paso con número, título, pasos y nota opcional."""
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(8)
    p_titulo.paragraph_format.space_after  = Pt(4)
    _run(p_titulo, f"  {numero}  ", bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=11)
    _run(p_titulo, f"  {titulo}", bold=True, color=C_TITULO, size=11)

    for paso in pasos:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, paso, size=10.5)

    if nota:
        p_nota = doc.add_paragraph()
        p_nota.paragraph_format.left_indent = Cm(0.5)
        p_nota.paragraph_format.space_after  = Pt(6)
        _run(p_nota, "💡 ", size=10.5)
        _run(p_nota, nota, italic=True, color=C_GRIS, size=10.5)

    if advertencia:
        p_adv = doc.add_paragraph()
        p_adv.paragraph_format.left_indent = Cm(0.5)
        p_adv.paragraph_format.space_after  = Pt(6)
        _run(p_adv, "⚠️ ", size=10.5)
        _run(p_adv, advertencia, bold=True, color=C_NARANJA, size=10.5)


def _tabla_nomenclatura(doc):
    """Tabla de nombres de archivos con colores."""
    headers = ["Depósito", "Nombre que genera Flexxus", "Nombre correcto a usar", "Atajo rápido"]
    rows = [
        ("SAN JOSE",        "Planilla de Stock_DD-MM-YYYY HH-MM-SS.xlsx", "SJ_stock_YYYY-MM-DD.xlsx",       "SJ Planilla de Stock_...xlsx"),
        ("LARREA",          "Planilla de Stock_DD-MM-YYYY HH-MM-SS.xlsx", "LAR_stock_YYYY-MM-DD.xlsx",      "LAR Planilla de Stock_...xlsx"),
        ("ES LOCAL / SARM.","Planilla de Stock_DD-MM-YYYY HH-MM-SS.xlsx", "ESLOCAL_stock_YYYY-MM-DD.xlsx",  "LAR Planilla de Stock_...xlsx"),
        ("FULL ML",         "Planilla de Stock_DD-MM-YYYY HH-MM-SS.xlsx", "FML_stock_YYYY-MM-DD.xlsx",      "FML Planilla de Stock_...xlsx"),
    ]
    col_w = [Cm(3.2), Cm(5.8), Cm(4.5), Cm(4.5)]
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header
    for i, (h, w) in enumerate(zip(headers, col_w)):
        cell = tbl.cell(0, i)
        cell.width = w
        _set_cell_bg(cell, "0A84FF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _run(p, h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=10)

    for r_idx, (dep, flexxus, correcto, atajo) in enumerate(rows, start=1):
        row_vals = [dep, flexxus, correcto, atajo]
        bg = "F2F2F7" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_vals):
            cell = tbl.cell(r_idx, c_idx)
            cell.width = col_w[c_idx]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            color = C_VERDE if c_idx == 2 else (C_NARANJA if c_idx == 1 else C_NEGRO)
            bold  = c_idx in (0, 2)
            _run(p, val, bold=bold, color=color, size=9.5)

    doc.add_paragraph()


def _tabla_otros_archivos(doc):
    """Tabla de otros archivos (sin renombrado)."""
    headers = ["Tipo de Archivo", "Módulo en Flexxus", "¿Renombrar?", "Filtros clave"]
    rows = [
        ("Optimización de Stock",  "Stock → Optimización de Stock",      "NO — nombre automático",  "Super Rubro: MODULOS · Período: 6 meses · Demanda: 30 días"),
        ("Lista de Precios",       "Ventas → Lista de Precios Editable",  "NO — nombre automático",  "Todas las listas · Moneda: USD · Incluir P.Comp"),
        ("Ventas por Artículo",    "Informes → Ventas por Artículo",      "NO — lleva 'Resumida'",   "Vista: Resumida · Rango: últimos 30–90 días · Todos los depósitos"),
        ("Compras por Marca",      "Informes → Compras por Marca",        "NO — lleva 'por Marca'",  "Activar: Facturas + NC + Remitos RE · Todas las marcas"),
        ("Cotización AI-TECH",     "Proveedor externo (Diego/FR)",        "NO — nombre libre",       "Debe contener 'cotizacion' o 'ai-tech' en el nombre"),
        ("Archivo Mariano",        "Interno (multi-hoja)",                "NO — lleva 'optimizacion'","NO debe contener 'stock' en el nombre"),
    ]
    col_w = [Cm(3.5), Cm(4.5), Cm(2.8), Cm(7.2)]
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (h, w) in enumerate(zip(headers, col_w)):
        cell = tbl.cell(0, i)
        cell.width = w
        _set_cell_bg(cell, "1C1C1E")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _run(p, h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=10)

    for r_idx, (tipo, modulo, renombrar, filtros) in enumerate(rows, start=1):
        bg = "F2F2F7" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate([tipo, modulo, renombrar, filtros]):
            cell = tbl.cell(r_idx, c_idx)
            cell.width = col_w[c_idx]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if c_idx == 2:
                ok = "NO" in val
                color = C_VERDE if ok else C_ROJO
            else:
                color = C_NEGRO
            _run(p, val, color=color, size=9.5, bold=(c_idx == 0))

    doc.add_paragraph()


def build():
    doc = Document()

    # ── Estilos globales ────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    # ── Márgenes ────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ══════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "ROKER NEXUS", bold=True, color=C_TITULO, size=28)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    _run(p2, "Guía Operativa del Sistema", bold=False, color=C_GRIS, size=16)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(40)
    _run(p3, "EL CELU — Versión 2.3.0 · Marzo 2026", italic=True, color=C_GRIS, size=11)

    # Recuadro de aviso
    box_p = doc.add_paragraph()
    box_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box_p.paragraph_format.space_after = Pt(60)
    _run(box_p, "⚠️  Este documento es de uso interno. No compartir externamente.", bold=True, color=C_NARANJA, size=10.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 1 — INTRODUCCIÓN
    # ══════════════════════════════════════════════════════════════
    _heading(doc, "1. ¿Qué es Roker Nexus?", level=1)
    _para(doc, "Roker Nexus es el sistema de gestión de inventario y análisis de datos de EL CELU. "
               "Permite visualizar el stock en tiempo real, analizar ventas, gestionar pedidos a proveedores "
               "y detectar anomalías — todo desde un navegador web.")
    _para(doc, "Acceso: abrís el link del sistema en cualquier navegador. No necesitás instalar nada.")

    _heading(doc, "Quién usa el sistema", level=2, color=C_NEGRO)
    p = _para(doc)
    _run(p, "• Roker (responsable principal):", bold=True, size=10.5)
    _run(p, " análisis, configuración, pedidos.", size=10.5)
    p2 = _para(doc)
    _run(p2, "• Operador delegado:", bold=True, size=10.5)
    _run(p2, " carga de archivos de Flexxus, consulta de stock.", size=10.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 2 — ARCHIVOS QUE NECESITA EL SISTEMA
    # ══════════════════════════════════════════════════════════════
    _heading(doc, "2. Qué archivos necesita el sistema", level=1)
    _para(doc, "El sistema se alimenta de archivos Excel que se exportan desde Flexxus. "
               "Hay 4 tipos de archivos principales que se cargan regularmente.")

    tipos = [
        ("📦 Planilla de Stock",     "URGENTE — Semanal o ante quiebre",     C_NARANJA),
        ("📊 Optimización de Stock", "Mensual",                               C_TITULO),
        ("💰 Lista de Precios",      "Al cambiar precios",                    C_VERDE),
        ("📈 Ventas / Compras",      "Mensual / al cierre del período",       C_GRIS),
    ]

    for emoji_titulo, frecuencia, color in tipos:
        p = _para(doc, indent=True)
        _run(p, emoji_titulo + " — ", bold=True, color=color, size=10.5)
        _run(p, frecuencia, color=C_GRIS, size=10.5)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 3 — PLANILLA DE STOCK (LA MÁS IMPORTANTE)
    # ══════════════════════════════════════════════════════════════
    _heading(doc, "3. Planilla de Stock — el archivo más crítico", level=1)

    # ALERTA PRINCIPAL
    p_alert = doc.add_paragraph()
    p_alert.paragraph_format.space_before = Pt(4)
    p_alert.paragraph_format.space_after  = Pt(12)
    _run(p_alert, "⚠️  IMPORTANTE: ", bold=True, color=C_NARANJA, size=11)
    _run(p_alert, "Flexxus siempre genera el mismo nombre de archivo para cualquier depósito. "
                  "Sin el prefijo correcto, el sistema no sabe de qué depósito es el stock y puede "
                  "sobreescribir datos incorrectamente.", color=C_NEGRO, size=10.5)

    _heading(doc, "3.1 Dónde descargarlo en Flexxus", level=2, color=C_NEGRO)
    p = _para(doc)
    _run(p, "Ruta exacta: ", bold=True, size=10.5)
    _run(p, "Stock  →  Listado General  →  [seleccionar depósito]  →  Exportar", color=C_TITULO, size=10.5)

    doc.add_paragraph()
    p_sc = doc.add_paragraph()
    p_sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_sc, "[ CAPTURA DE PANTALLA — Flexxus: Stock > Listado General ]",
         italic=True, color=C_GRIS, size=10)
    p_sc.paragraph_format.space_after = Pt(2)
    p_sc2 = doc.add_paragraph()
    p_sc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_sc2, "Reemplazar este bloque con una captura real de Flexxus",
         italic=True, color=RGBColor(0xAE,0xAE,0xB2), size=9)
    p_sc2.paragraph_format.space_after = Pt(14)

    _heading(doc, "3.2 Pasos para exportar — SAN JOSE", level=2, color=C_NEGRO)
    _step_box(doc, "1", "Ir al módulo de Stock",
              ["En Flexxus, hacer clic en el menú lateral → Stock",
               "Seleccionar 'Listado General' (no 'Optimización')"])

    _step_box(doc, "2", "Seleccionar el depósito",
              ["En el filtro 'Depósito', seleccionar únicamente: SAN JOSE",
               "NO seleccionar 'Todos' — exportar de a UN depósito por vez"],
              advertencia="Si seleccionás 'Todos' el sistema no puede separar el stock por depósito")

    _step_box(doc, "3", "Configurar filtros",
              ["Sin filtro de fecha — muestra el stock actual",
               "Sin filtro de rubro — descargar TODOS los artículos del depósito",
               "Sin filtro de marca"],
              nota="Filtrar por rubro MODULOS haría que se pierdan todos los demás artículos")

    _step_box(doc, "4", "Exportar",
              ["Clic en el botón Exportar / Excel",
               "El archivo se descarga automáticamente como:",
               "  Planilla de Stock_15-03-2026 10-22-45.xlsx"])

    _step_box(doc, "5", "Renombrar el archivo ANTES de subirlo",
              ["Hacer clic derecho sobre el archivo → Cambiar nombre",
               "Poner el nombre en este formato exacto:   SJ_stock_2026-03-15.xlsx",
               "Atajo rápido: también funciona poner   SJ   al INICIO del nombre original"],
              nota="El sistema reconoce el depósito por el prefijo del nombre. Sin prefijo = error.")

    _step_box(doc, "6", "Subir al sistema",
              ["Abrir Roker Nexus en el navegador",
               "Ir al menú lateral → ⚙️ SISTEMA → 📥 Importar",
               "Arrastrar el archivo a la zona de carga (tab: Carga automática)",
               "El sistema muestra: ✅ X filas importadas"])

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # TABLA DE NOMENCLATURA
    # ══════════════════════════════════════════════════════════════
    _heading(doc, "3.3 Tabla de nomenclatura de archivos de stock", level=2, color=C_NEGRO)
    _para(doc, "Esta tabla resume los nombres exactos para cada depósito:")
    _tabla_nomenclatura(doc)

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 4 — OTROS ARCHIVOS
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, "4. Otros archivos del sistema", level=1)

    _heading(doc, "4.1 Optimización de Stock", level=2, color=C_NEGRO)
    p = _para(doc)
    _run(p, "Ruta: ", bold=True, size=10.5)
    _run(p, "Stock → Optimización de Stock", color=C_TITULO, size=10.5)
    _step_box(doc, "1", "Configurar en Flexxus",
              ["Super Rubro: seleccionar MODULOS",
               "Período de análisis: 6 meses",
               "Demanda promedio: 30 días",
               "Sin filtro de depósito (analiza stock combinado)"],
              nota="Este archivo NO necesita renombrarse. El nombre que genera Flexxus ya es reconocido automáticamente.")

    _heading(doc, "4.2 Lista de Precios", level=2, color=C_NEGRO)
    p = _para(doc)
    _run(p, "Ruta: ", bold=True, size=10.5)
    _run(p, "Ventas → Lista de Precios Editable", color=C_TITULO, size=10.5)
    _step_box(doc, "1", "Configurar en Flexxus",
              ["Sin filtro de artículo — exportar TODA la lista",
               "Incluir todas las listas: 1 (mayorista), 2, 3, 4 (ML), 5",
               "Moneda: USD (dólares) — NO pesos",
               "Incluir columna P.Comp (precio de compra)"],
              advertencia="Si exportás en ARS en vez de USD, todos los precios van a estar mal en el sistema")

    _heading(doc, "4.3 Ventas por Artículo", level=2, color=C_NEGRO)
    p = _para(doc)
    _run(p, "Ruta: ", bold=True, size=10.5)
    _run(p, "Informes → Ventas por Artículo (vista Resumida)", color=C_TITULO, size=10.5)
    _step_box(doc, "1", "Configurar en Flexxus",
              ["Usar la vista 'Resumida' (no la detallada)",
               "Rango de fechas: últimos 30 a 90 días",
               "Sin filtro de artículo ni depósito — todas las ventas",
               "Exportar a Excel"],
              nota="El nombre con 'Resumida' es reconocido automáticamente. Para análisis de 6 meses, rango desde el 01/10/2025.")

    _heading(doc, "4.4 Tabla resumen — todos los archivos", level=2, color=C_NEGRO)
    _tabla_otros_archivos(doc)

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 5 — CÓMO SUBIR ARCHIVOS AL SISTEMA
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, "5. Cómo subir archivos al sistema", level=1)
    _para(doc, "Una vez que tenés los archivos en tu computadora con el nombre correcto, "
               "los subís al sistema desde la página de Importar.")

    _step_box(doc, "1", "Abrir la página de Importar",
              ["En el menú lateral izquierdo, ir a: ⚙️ SISTEMA → 📥 Importar",
               "O hacer clic en la sección Sistema del menú"])

    p_sc3 = doc.add_paragraph()
    p_sc3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sc3.paragraph_format.space_after = Pt(2)
    _run(p_sc3, "[ CAPTURA DE PANTALLA — Roker Nexus: página Importar ]",
         italic=True, color=C_GRIS, size=10)
    p_sc4 = doc.add_paragraph()
    p_sc4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sc4.paragraph_format.space_after = Pt(14)
    _run(p_sc4, "Reemplazar con captura real del sistema",
         italic=True, color=RGBColor(0xAE,0xAE,0xB2), size=9)

    _step_box(doc, "2", "Usar la tab 'Carga automática'",
              ["Está seleccionada por defecto",
               "Arrastrar el archivo a la zona de carga, o hacer clic para seleccionarlo",
               "Se pueden subir varios archivos a la vez"])

    _step_box(doc, "3", "Verificar que el sistema lo reconoció",
              ["El sistema muestra el tipo de archivo detectado",
               "Si dice ✅ X filas importadas: todo OK",
               "Si dice ⚠️ Tipo no reconocido: el archivo necesita renombrarse"],
              advertencia="Si el archivo no es reconocido, el sistema te ofrece un selector para indicar manualmente qué tipo es.")

    _step_box(doc, "4", "Usar el Renombrador rápido si hay dudas",
              ["Ir a la tab '📖 Guía de archivos'",
               "Seleccionar el depósito en el desplegable",
               "Pegar el nombre original del archivo",
               "El sistema te dice exactamente cómo renombrarlo"])

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 6 — PREGUNTAS FRECUENTES
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, "6. Preguntas frecuentes", level=1)

    faqs = [
        ("¿Qué pasa si subo el mismo archivo dos veces?",
         "El sistema lo sobreescribe con los datos nuevos. No hay duplicados. Si fue un error, "
         "podés volver a subir la versión anterior."),
        ("¿Cada cuánto hay que cargar la Planilla de Stock?",
         "Idealmente una vez por semana, o siempre que haya quiebres importantes que revisar. "
         "El sistema muestra cuántos días pasaron desde la última carga en el panel de salud de datos."),
        ("¿Qué pasa si me olvido de renombrar el archivo de stock?",
         "El sistema te avisa con ⚠️ Tipo no reconocido. Podés seleccionar el tipo manualmente "
         "y el depósito desde el menú que aparece, sin necesidad de cerrar y renombrar."),
        ("El sistema dice 'Sin importador', ¿qué hago?",
         "Ese tipo de archivo no tiene un importador configurado todavía. Contactar a Roker."),
        ("¿Necesito subir los archivos en un orden específico?",
         "No hay orden obligatorio, pero lo más útil es empezar por Stock (para tener datos actualizados) "
         "y después Optimización y Precios."),
        ("¿Puedo subir archivos desde el celular?",
         "Sí, el sistema es accesible desde el navegador del celular, aunque la experiencia es "
         "mejor desde computadora."),
    ]

    for pregunta, respuesta in faqs:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after  = Pt(2)
        _run(p_q, "❓ " + pregunta, bold=True, color=C_TITULO, size=10.5)
        p_a = _para(doc, indent=True, space_after=4)
        _run(p_a, respuesta, size=10.5, color=C_NEGRO)

    # ══════════════════════════════════════════════════════════════
    # SECCIÓN 7 — CONTACTO
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, "7. Contacto y soporte", level=1)
    _para(doc, "Para cualquier problema con el sistema, contactar a Roker.")
    _para(doc, "No modificar configuraciones del sistema sin consultar.")

    p_roker = _para(doc, indent=True)
    _run(p_roker, "Roker: ", bold=True, size=10.5, color=C_TITULO)
    _run(p_roker, "kuaji28@gmail.com", size=10.5)

    doc.add_paragraph()
    p_ver = _para(doc)
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_ver, "Roker Nexus v2.3.0 · EL CELU · Marzo 2026",
         italic=True, color=C_GRIS, size=9.5)

    # ── Guardar ─────────────────────────────────────────────────────
    out_path = "guia_operativa_nexus.docx"
    doc.save(out_path)
    print(f"✅ Guardado: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
